import os
import sys
import json
import time
import math
import logging
import threading
import queue
import hashlib
from dataclasses import dataclass
from pathlib import Path
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
from datetime import datetime, timezone
import atexit

# --- GUI Imports ---
import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox

# --- Charting Imports ---
try:
    import matplotlib
    matplotlib.use("TkAgg")
    from matplotlib.figure import Figure
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False


# --- Original third-party imports (best-effort) ---
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.prompts import PromptTemplate
    from langchain.docstore.document import Document
    from langchain_community.vectorstores import FAISS
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_google_vertexai import VertexAIEmbeddings, ChatVertexAI
    from vertexai.language_models import TextEmbeddingModel
except Exception:
    RecursiveCharacterTextSplitter, PromptTemplate, Document, FAISS = None, None, None, None
    HuggingFaceEmbeddings, VertexAIEmbeddings, ChatVertexAI, TextEmbeddingModel = None, None, None, None

try:
    import tiktoken
except Exception:
    tiktoken = None

try:
    from transformers import AutoTokenizer
except Exception:
    AutoTokenizer = None

try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None

try:
    import docx
except Exception:
    docx = None

try:
    from PIL import Image, ImageEnhance
    import pytesseract
except Exception:
    Image, pytesseract = None, None

try:
    from langdetect import detect, DetectorFactory
    DetectorFactory.seed = 0
except Exception:
    def detect(text): return "en"

# ---------- CONFIG ----------
@dataclass
class AppConfig:
    GOOGLE_APPLICATION_CREDENTIALS: str = r"D:\GEN_AI_25\gemini-api-key.json" # <-- IMPORTANT: UPDATE THIS PATH 
    DATA_DIR: str = "data"
    CACHE_DIR: str = "cache"
    MEMORY_DIR: str = "memory"
    PROFILE_PATH: str = "profile.json"
    STORES_DIR: str = "stores"
    CHUNK_SIZE: int = 600
    CHUNK_OVERLAP: int = 60
    MIN_CHUNK_LENGTH: int = 50
    MAX_CHUNKS: int = 1000
    TOP_K: int = 5
    DEDUP_THRESHOLD: float = 0.78
    MODEL_NAME: str = "gemini-2.5-flash-lite" 
    TEMPERATURE: float = 0.15
    MAX_OUTPUT_TOKENS: int = 1024
    MEMORY_TOP_K: int = 3
    MEMORY_TRUNCATE_CHARS: int = 420
    CONTEXT_TRUNCATE_CHARS: int = 900
    FAISS_SAVE_EVERY: int = 8
    SUMMARIZE_BATCH: int = 8
    SLIDING_WINDOW: int = 5
    INDEX_THREADS: int = 2

    def __post_init__(self):
        Path(self.DATA_DIR).mkdir(parents=True, exist_ok=True)
        Path(self.CACHE_DIR).mkdir(parents=True, exist_ok=True)
        Path(self.MEMORY_DIR).mkdir(parents=True, exist_ok=True)
        Path(self.STORES_DIR).mkdir(parents=True, exist_ok=True)

CONFIG = AppConfig()
if CONFIG.GOOGLE_APPLICATION_CREDENTIALS and Path(CONFIG.GOOGLE_APPLICATION_CREDENTIALS).exists():
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CONFIG.GOOGLE_APPLICATION_CREDENTIALS

# ---------- Logging ----------
logfile = Path(CONFIG.DATA_DIR) / f"app_{datetime.now():%Y%m%d_%H%M%S}.log"
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
    handlers=[logging.FileHandler(logfile, encoding="utf-8"), logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger("legal-ai")

def tail_log_lines(path: Path, n: int = 200) -> str:
    try:
        with open(path, "r", encoding="utf-8") as f:
            lines = f.readlines()
        return "".join(lines[-n:])
    except Exception:
        return ""

# ---------- Utilities ----------
def file_fingerprint(path: str) -> str:
    try:
        st = Path(path).stat()
        s = f"{path}|{st.st_mtime_ns}|{st.st_size}"
        return hashlib.md5(s.encode()).hexdigest()
    except Exception:
        return hashlib.md5(path.encode()).hexdigest()

def folder_size_mb(path: Path) -> float:
    total = 0
    try:
        for entry in os.scandir(path):
            if entry.is_file():
                total += entry.stat().st_size
            elif entry.is_dir():
                total += folder_size_mb(Path(entry.path)) * 1024 * 1024
    except Exception:
        pass
    return total / (1024 * 1024)

def compact_profile_for_prompt(profile: Dict[str, Any]) -> Dict[str, str]:
    return {k: str(v) for k, v in profile.items() if v and k != "_extra"}

# ---------- Token monitor ----------
class TokenMonitor:
    def __init__(self):
        self.enc = None
        if tiktoken:
            try:
                self.enc = tiktoken.get_encoding("cl100k_base")
            except Exception:
                self.enc = None
        try:
            self.bert_tok = AutoTokenizer.from_pretrained("nlpaueb/legal-bert-base-uncased") if AutoTokenizer else None
        except Exception:
            self.bert_tok = None

    def count_llm_tokens(self, text: str) -> int:
        if not text: return 0
        if self.enc:
            try: return len(self.enc.encode(text))
            except Exception: pass
        return max(1, int(len(text.split()) * 1.3))

    def count_bert_tokens(self, text: str) -> int:
        if not text: return 0
        if self.bert_tok:
            try: return len(self.bert_tok.encode(text, add_special_tokens=False))
            except Exception: pass
        return max(1, int(len(text.split()) * 1.0))

# ---------- Extractor (document & image) ----------
class Extractor:
    def __init__(self, cache_dir: Path):
        self.cache_dir = cache_dir
        self.cache_dir.mkdir(parents=True, exist_ok=True)

    def from_pdf(self, pdf_path: str) -> str:
        fid = file_fingerprint(pdf_path)
        cpath = self.cache_dir / f"extract_{fid}.txt"
        if cpath.exists(): return cpath.read_text(encoding="utf-8")
        text_parts = []
        try:
            if PdfReader is None: raise RuntimeError("PyPDF2 not installed")
            with open(pdf_path, "rb") as f:
                reader = PdfReader(f)
                for i, page in enumerate(reader.pages, start=1):
                    txt = (page.extract_text() or "").strip()
                    if txt: text_parts.append(f"\n--- Page {i} ---\n{txt}")
        except Exception: logger.exception("Error reading PDF %s", pdf_path)
        text = "\n".join(text_parts).strip() if text_parts else "No readable text found."
        try: cpath.write_text(text, encoding="utf-8")
        except Exception: logger.exception("Failed to cache PDF extraction %s", cpath)
        return text

    def from_docx(self, docx_path: str) -> str:
        fid = file_fingerprint(docx_path)
        cpath = self.cache_dir / f"extract_{fid}.txt"
        if cpath.exists(): return cpath.read_text(encoding="utf-8")
        parts = []
        try:
            if docx is None: raise RuntimeError("python-docx not installed")
            doc = docx.Document(docx_path)
            for para in doc.paragraphs:
                txt = para.text.strip()
                if txt: parts.append(txt)
        except Exception: logger.exception("DOCX read error %s", docx_path)
        text = "\n".join(parts).strip() if parts else "No readable text found."
        try: cpath.write_text(text, encoding="utf-8")
        except Exception: logger.exception("Failed to cache DOCX extraction %s", cpath)
        return text

    def from_text(self, txt_path: str) -> str:
        try: return Path(txt_path).read_text(encoding="utf-8")
        except Exception:
            logger.exception("Error reading text file %s", txt_path)
            return ""

    def from_json(self, json_path: str) -> str:
        try:
            data = json.loads(Path(json_path).read_text(encoding="utf-8"))
            return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception:
            logger.exception("Error reading json file %s", json_path)
            return ""

    def from_image(self, image_path: str, lang: str = "eng") -> str:
        fid = file_fingerprint(image_path)
        cpath = self.cache_dir / f"extract_{fid}.txt"
        if cpath.exists(): return cpath.read_text(encoding="utf-8")
        text = ""
        try:
            if Image is None or pytesseract is None: raise RuntimeError("Pillow or pytesseract not installed")
            with Image.open(image_path) as im:
                if im.mode != "L": im = im.convert("L")
                im = ImageEnhance.Contrast(im).enhance(1.6)
                im = ImageEnhance.Sharpness(im).enhance(1.5)
                text = pytesseract.image_to_string(im, config="--oem 3 --psm 6", lang=lang) or ""
        except Exception: logger.exception("Image OCR error for %s", image_path)
        if not text.strip(): text = "No readable text found."
        try: cpath.write_text(text, encoding="utf-8")
        except Exception: logger.exception("Failed to cache image extraction %s", cpath)
        return text

# ---------- Dual Vector Store ----------
class DualVectorStore:
    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self.base_dir.mkdir(parents=True, exist_ok=True)
        self.vertex_emb, self.legal_emb = None, None
        self.vertex_store: Optional[FAISS] = None
        self.legal_store: Optional[FAISS] = None
        self._lock = threading.RLock()
        self._vertex_add_count, self._legal_add_count = 0, 0
        self._init_embeddings()

    def _vs_path(self, fid: str, kind: str) -> Path: return self.base_dir / f"vs_{kind}_{fid}"

    def _init_embeddings(self):
        try:
            if VertexAIEmbeddings is None: raise RuntimeError("VertexAIEmbeddings not available")
            self.vertex_emb = VertexAIEmbeddings(model_name="text-embedding-004")
            logger.info("VertexAIEmbeddings initialized.")
        except Exception: logger.exception("VertexAIEmbeddings init failed.")
        try:
            if HuggingFaceEmbeddings is None: raise RuntimeError("HuggingFaceEmbeddings not available")
            self.legal_emb = HuggingFaceEmbeddings(model_name="nlpaueb/legal-bert-base-uncased", model_kwargs={"device": "cpu"}, encode_kwargs={"normalize_embeddings": True})
            logger.info("HuggingFace legal-bert embeddings initialized.")
        except Exception: logger.exception("HuggingFaceEmbeddings init failed.")

    def _load_or_create_store(self, fid: str):
        with self._lock:
            if self.vertex_emb:
                vpath = self._vs_path(fid, "vertex")
                try:
                    if vpath.exists(): self.vertex_store = FAISS.load_local(vpath.as_posix(), self.vertex_emb, allow_dangerous_deserialization=True)
                    else:
                        self.vertex_store = FAISS.from_documents([Document(page_content="seed")], self.vertex_emb)
                        self.vertex_store.save_local(vpath.as_posix())
                except Exception: logger.exception(f"Vertex store error for {fid}")

            if self.legal_emb:
                lpath = self._vs_path(fid, "legal")
                try:
                    if lpath.exists(): self.legal_store = FAISS.load_local(lpath.as_posix(), self.legal_emb, allow_dangerous_deserialization=True)
                    else:
                        self.legal_store = FAISS.from_documents([Document(page_content="seed")], self.legal_emb)
                        self.legal_store.save_local(lpath.as_posix())
                except Exception: logger.exception(f"Legal store error for {fid}")

    def build_or_load(self, text: Optional[str], fid: str):
        if RecursiveCharacterTextSplitter is None:
            docs: List[Document] = [Document(page_content=text or "")]
        else:
            splitter = RecursiveCharacterTextSplitter(chunk_size=CONFIG.CHUNK_SIZE, chunk_overlap=CONFIG.CHUNK_OVERLAP)
            docs = splitter.create_documents([text]) if text else []
            docs = [d for d in docs if len(d.page_content.strip()) >= CONFIG.MIN_CHUNK_LENGTH]
            if len(docs) > CONFIG.MAX_CHUNKS: docs = docs[:CONFIG.MAX_CHUNKS]

        with self._lock:
            if self.vertex_emb:
                v_path = self._vs_path(fid, "vertex")
                try:
                    self.vertex_store = FAISS.from_documents(docs or [Document(page_content="seed")], self.vertex_emb)
                    self.vertex_store.save_local(v_path.as_posix())
                    logger.info("Created/updated vertex store %s", v_path.name)
                except Exception: logger.exception("Vertex store build error for %s", fid)

            if self.legal_emb:
                l_path = self._vs_path(fid, "legal")
                try:
                    self.legal_store = FAISS.from_documents(docs or [Document(page_content="seed")], self.legal_emb)
                    self.legal_store.save_local(l_path.as_posix())
                    logger.info("Created/updated legal store %s", l_path.name)
                except Exception: logger.exception("Legal store build error for %s", fid)

    def add_doc_to_stores(self, doc: Document, fid: str):
        with self._lock:
            if self.vertex_store:
                try:
                    self.vertex_store.add_documents([doc])
                    self._vertex_add_count += 1
                    if self._vertex_add_count % CONFIG.FAISS_SAVE_EVERY == 0:
                        self.vertex_store.save_local(self._vs_path(fid, "vertex").as_posix())
                except Exception: logger.exception("Error adding to vertex store")
            if self.legal_store:
                try:
                    self.legal_store.add_documents([doc])
                    self._legal_add_count += 1
                    if self._legal_add_count % CONFIG.FAISS_SAVE_EVERY == 0:
                        self.legal_store.save_local(self._vs_path(fid, "legal").as_posix())
                except Exception: logger.exception("Error adding to legal store")

    def save_all(self, fid: str):
        with self._lock:
            try:
                if self.vertex_store: self.vertex_store.save_local(self._vs_path(fid, "vertex").as_posix())
                if self.legal_store: self.legal_store.save_local(self._vs_path(fid, "legal").as_posix())
            except Exception: logger.exception("Error saving vectorstores")

    def search(self, query: str, k: int) -> List[str]:
        results: List[Tuple[str, float]] = []
        with self._lock:
            if self.vertex_store:
                try:
                    vs = self.vertex_store.similarity_search_with_score(query, k=max(1, k * 3))
                    results += [(doc.page_content, float(score)) for doc, score in vs]
                except Exception: logger.exception("Vertex search error")
            if self.legal_store:
                try:
                    ls = self.legal_store.similarity_search_with_score(query, k=max(1, k * 3))
                    results += [(doc.page_content, float(score)) for doc, score in ls]
                except Exception: logger.exception("Legal search error")
        
        def jaccard(a: str, b: str) -> float:
            A, B = set(a.lower().split()), set(b.lower().split())
            return len(A & B) / max(1, len(A | B))

        seen_texts: List[str] = []
        unique_results: List[str] = []
        for text, score in sorted(results, key=lambda x: -x[1]):
            if any(jaccard(text, s) >= CONFIG.DEDUP_THRESHOLD for s in seen_texts): continue
            seen_texts.append(text)
            unique_results.append(text)
            if len(unique_results) >= k: break
        return unique_results

# ---------- ProfileMemory ----------
class ProfileMemory:
    SAFE_KEYS = {"name", "age", "gender", "location", "profession", "organization", "interests", "email", "phone", "language"}
    def __init__(self, path: Path):
        self.path = path
        if not self.path.exists(): self._write({})
        self._data = self._read()
        self._lock = threading.RLock()

    def _read(self) -> Dict[str, Any]:
        try: return json.loads(self.path.read_text(encoding="utf-8"))
        except Exception:
            logger.exception("Failed to read profile.json"); return {}

    def _write(self, data: Dict[str, Any]):
        try: self.path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception: logger.exception("Failed to write profile.json")

    def set(self, key: str, value: Any):
        key = key.strip()
        with self._lock:
            if key in self.SAFE_KEYS: self._data[key] = value
            else: self._data.setdefault("_extra", {})[key] = value
            self._write(self._data)

    def update_from_dict(self, d: Dict[str, Any]):
        if not isinstance(d, dict): return
        with self._lock:
            for k, v in d.items():
                if k in self.SAFE_KEYS: self._data[k] = v
                else: self._data.setdefault("_extra", {})[k] = v
            self._write(self._data)

    def get(self, key: str, default: Optional[Any] = None) -> Optional[Any]: return self._data.get(key, default)
    def all(self) -> Dict[str, Any]: return dict(self._data)

# ---------- MemoryStore ----------
class MemoryStore:
    def __init__(self, mem_dir: Path, vs: DualVectorStore, extractor_llm=None):
        self.mem_dir = mem_dir
        self.vs = vs
        self.fid = "chat_memory"
        self.buffer: List[Dict[str, Any]] = []
        self.buffer_limit = CONFIG.SLIDING_WINDOW
        self.summarize_batch = CONFIG.SUMMARIZE_BATCH
        self._summ_queue: "queue.Queue[List[Dict]]" = queue.Queue()
        self._summ_worker = threading.Thread(target=self._summarization_worker, daemon=True)
        self.extractor_llm = extractor_llm
        try:
            self.vs._load_or_create_store(self.fid)
            self._summ_worker.start()
        except Exception: logger.exception("Failed to initialize memory vector stores")

    def _truncated(self, role: str, content: str) -> str: return f"{role}: {content[:CONFIG.MEMORY_TRUNCATE_CHARS]}"

    def add(self, role: str, content: str):
        entry = {"role": role, "text": content, "ts": datetime.now(timezone.utc).isoformat()}
        self.buffer.append(entry)
        short = self._truncated(role, content)
        try: self.vs.add_doc_to_stores(Document(page_content=short, metadata={"type": "mem_truncate"}), self.fid)
        except Exception: logger.exception("Error adding truncated memory to vectorstores")
        if len(self.buffer) > self.buffer_limit:
            n = min(self.summarize_batch, len(self.buffer) - self.buffer_limit)
            batch, self.buffer = self.buffer[:n], self.buffer[n:]
            self._summ_queue.put(batch)

    def _summarization_worker(self):
        while True:
            try:
                batch = self._summ_queue.get()
                if batch is None: break
                conv_blob = "\n".join([f"{e['role']}: {e['text']}" for e in batch])
                summary = self._summarize_conv(conv_blob)
                self.vs.add_doc_to_stores(Document(page_content=f"MEMORY_SUMMARY: {summary}", metadata={"type": "memory_summary"}), self.fid)
                self.vs.save_all(self.fid)
                self._summ_queue.task_done()
            except Exception: logger.exception("Summarization worker error"); time.sleep(1)

    def _summarize_conv(self, text: str) -> str:
        if not text.strip(): return "No content."
        try:
            if self.extractor_llm:
                prompt = f"Summarize the following conversation into one short, factual sentence. Do NOT include personal identifiers.\n\n{text}\n\nSummary:"
                out = self.extractor_llm.invoke(prompt)
                return out.content.strip() if hasattr(out, "content") else str(out).strip()[:400]
        except Exception: logger.exception("Vertex summarization failed")
        return (" ; ".join(text.splitlines()[:2]))[:400]

    def retrieve(self, query: str, k: int, slider_k: int) -> List[str]:
        results = []
        try: results = self.vs.search(query, k=k)
        except Exception: logger.exception("MemoryStore retrieve error")
        recent_count = max(0, min(len(self.buffer), slider_k))
        recent_contexts = [f"{e['role']}: {e['text']}" for e in self.buffer[-recent_count:]]
        merged = recent_contexts + results
        seen = set()
        final = []
        for item in merged:
            key = item.strip()[:500]
            if key not in seen: final.append(item); seen.add(key)
            if len(final) >= k: break
        return final

    def shutdown(self):
        try: self._summ_queue.put(None); self._summ_worker.join(timeout=5)
        except Exception: pass

# ---------- LLM-based Profile Extractor ----------
class ProfileExtractor:
    def __init__(self, llm_client=None): self.llm = llm_client

    def extract(self, text: str) -> Dict[str, Any]:
        text = text.strip()
        if not text: return {}
        if self.llm:
            try:
                prompt = f"Extract user profile attributes from the message. Return ONLY valid JSON. Use keys: name, age, location, profession, interests.\n\nMessage:\n{text}\n\nJSON:"
                out = self.llm.invoke(prompt)
                raw = out.content.strip() if hasattr(out, "content") else str(out).strip()
                start = raw.find("{"); end = raw.rfind("}") + 1
                return json.loads(raw[start:end]) if start != -1 and end != 0 else {}
            except Exception: logger.exception("LLM extraction failed, falling back to heuristic")
        return {} # Simple fallback

# ---------- QA Engine ----------
class QAEngine:
    def __init__(self, profile: ProfileMemory, vs: DualVectorStore, memory: MemoryStore, llm_client=None):
        self.profile, self.vs, self.memory, self.llm = profile, vs, memory, llm_client
        if PromptTemplate:
            self.prompt = PromptTemplate(input_variables=["context", "history", "profile", "question"], template="You are a concise, accurate assistant.\nPROFILE:\n{profile}\n\nRECENT:\n{history}\n\nCONTEXT:\n{context}\n\nQUESTION:\n{question}\n\nAnswer clearly.")
        else:
            self.prompt = "You are a concise, accurate assistant.\nPROFILE:\n{profile}\n\nRECENT:\n{history}\n\nCONTEXT:\n{context}\n\nQUESTION:\n{question}\n\nAnswer clearly."

    def answer(self, question: str, docs: List[str], memory_slider_val: int) -> Tuple[str, Dict[str, Any]]:
        profile_json = json.dumps(compact_profile_for_prompt(self.profile.all()), ensure_ascii=False)
        mem_results = self.memory.retrieve(question, k=CONFIG.MEMORY_TOP_K, slider_k=memory_slider_val)
        history = "\n".join(mem_results[:CONFIG.MEMORY_TOP_K])
        context = "\n".join([d[:CONFIG.CONTEXT_TRUNCATE_CHARS] for d in docs]) if docs else "No relevant documents."
        
        if PromptTemplate:
            filled = self.prompt.format(context=context, history=history, profile=profile_json, question=question)
        else:
            filled = self.prompt.replace("{context}", context).replace("{history}", history).replace("{profile}", profile_json).replace("{question}", question)

        metrics = {"llm_tokens": TOKEN_MON.count_llm_tokens(filled), "emb_tokens": TOKEN_MON.count_bert_tokens(question)}
        if not self.llm: return "LLM not initialized.", metrics
        t0 = time.time()
        try:
            out = self.llm.invoke(filled)
            metrics["latency"] = time.time() - t0
            return (out.content.strip() if hasattr(out, "content") else str(out).strip()), metrics
        except Exception as e:
            metrics["latency"] = time.time() - t0
            logger.exception("LLM call failed")
            return f"LLM invocation failed: {e}", metrics

# ---------- Indexing Manager ----------
class IndexingManager:
    def __init__(self, vs: DualVectorStore):
        self.vs = vs
        self._index_queue = queue.Queue()
        self._workers = [threading.Thread(target=self._worker, daemon=True) for _ in range(CONFIG.INDEX_THREADS)]
        for w in self._workers: w.start()

    def enqueue(self, text: str, fid: str): self._index_queue.put((text, fid))

    def _worker(self):
        while True:
            try:
                text, fid = self._index_queue.get()
                if fid is None: break
                logger.info(f"Index worker processing {fid}")
                self.vs.build_or_load(text, fid)
                logger.info(f"Index worker finished {fid}")
                self._index_queue.task_done()
            except Exception: logger.exception("Index worker error")

    def shutdown(self):
        for _ in self._workers: self._index_queue.put((None, None))
        for w in self._workers: w.join(timeout=2)

# ---------- Health Check ----------
def check_health(llm_client, vertex_emb, legal_emb, vs: DualVectorStore) -> Dict[str, str]:
    status = {}
    try:
        if llm_client: llm_client.invoke("ping"); status["LLM"] = "✅ OK"
        else: status["LLM"] = "❌ Not initialized"
    except Exception as e: status["LLM"] = f"❌ ERROR: {e}"
    try:
        if vertex_emb: vertex_emb.embed_query("health"); status["VertexAI Embeddings"] = "✅ OK"
        else: status["VertexAI Embeddings"] = "❌ Disabled"
    except Exception as e: status["VertexAI Embeddings"] = f"❌ ERROR: {e}"
    try:
        if legal_emb: legal_emb.embed_query("health"); status["LegalBERT Embeddings"] = "✅ OK"
        else: status["LegalBERT Embeddings"] = "❌ Disabled"
    except Exception as e: status["LegalBERT Embeddings"] = f"❌ ERROR: {e}"
    try:
        vs.search("health", k=1); status["FAISS Stores"] = "✅ OK"
    except Exception as e: status["FAISS Stores"] = f"❌ ERROR: {e}"
    try: status["StoresSizeMB"] = f"{folder_size_mb(Path(CONFIG.STORES_DIR)):.2f} MB"
    except Exception: pass
    logger.info(f"Health: {status}")
    return status

# ==============================================================================
# SECTION 3: TKINTER GUI APPLICATION
# ==============================================================================

class AppState:
    """A class to hold the application's state, replacing st.session_state."""
    def __init__(self):
        self.vs: Optional[DualVectorStore] = None
        self.profile: Optional[ProfileMemory] = None
        self.extractor_llm: Optional[ChatVertexAI] = None
        self.profile_extractor: Optional[ProfileExtractor] = None
        self.memory: Optional[MemoryStore] = None
        self.indexer: Optional[IndexingManager] = None
        self.qa: Optional[QAEngine] = None
        self.metrics = {
            "tokens": [], "latencies": [], "buffer_sizes": [],
            "store_sizes": [], "timestamps": []
        }
        self.chat_history: List[Tuple[str, str]] = []
        self.last_health: Dict[str, str] = {}
        self.memory_slider_val: int = CONFIG.SLIDING_WINDOW

class LegalAIGUI(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Legal AI Assistant")
        self.geometry("1400x900")

        self.state = AppState()
        self.response_queue = queue.Queue()

        self.init_backend()
        self.create_widgets()

        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.update_log_display() # Start periodic log updates

    def init_backend(self):
        """Initialize all backend components."""
        logger.info("Initializing backend components...")
        self.state.vs = DualVectorStore(Path(CONFIG.STORES_DIR))
        self.state.profile = ProfileMemory(Path(CONFIG.PROFILE_PATH))
        
        if ChatVertexAI:
            try:
                self.state.extractor_llm = ChatVertexAI(model=CONFIG.MODEL_NAME, temperature=0.0, max_output_tokens=512)
                qa_llm = ChatVertexAI(model=CONFIG.MODEL_NAME, temperature=CONFIG.TEMPERATURE, max_output_tokens=CONFIG.MAX_OUTPUT_TOKENS)
                logger.info("VertexAI LLMs initialized.")
            except Exception as e:
                logger.exception("LLM init failed.")
                messagebox.showerror("LLM Error", f"Failed to initialize VertexAI LLMs: {e}\n\nPlease check your GOOGLE_APPLICATION_CREDENTIALS path and authentication.")
                self.state.extractor_llm = None
                qa_llm = None
        else:
            logger.error("ChatVertexAI not available from langchain.")
            qa_llm = None
        
        self.state.profile_extractor = ProfileExtractor(llm_client=self.state.extractor_llm)
        self.state.memory = MemoryStore(Path(CONFIG.MEMORY_DIR), vs=self.state.vs, extractor_llm=self.state.extractor_llm)
        self.state.indexer = IndexingManager(self.state.vs)
        self.state.qa = QAEngine(self.state.profile, self.state.vs, self.state.memory, llm_client=qa_llm)
        logger.info("Backend initialization complete.")
        
        # Register cleanup functions
        atexit.register(self.on_closing)

    def create_widgets(self):
        """Create and layout all GUI widgets."""
        # Main layout: Paned window for resizable columns
        main_pane = ttk.PanedWindow(self, orient=tk.HORIZONTAL)
        main_pane.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # --- Sidebar (Left) ---
        sidebar_frame = ttk.Frame(main_pane, width=350)
        main_pane.add(sidebar_frame, weight=1)
        self._create_sidebar_widgets(sidebar_frame)

        # --- Main content (Right): Another paned window for chat and dashboard ---
        content_pane = ttk.PanedWindow(main_pane, orient=tk.HORIZONTAL)
        main_pane.add(content_pane, weight=4)

        chat_frame = ttk.Frame(content_pane, width=600)
        dashboard_frame = ttk.Frame(content_pane, width=450)
        content_pane.add(chat_frame, weight=2)
        content_pane.add(dashboard_frame, weight=1)

        self._create_chat_widgets(chat_frame)
        self._create_dashboard_widgets(dashboard_frame)

    def _create_sidebar_widgets(self, parent):
        parent.grid_rowconfigure(5, weight=1)
        parent.grid_columnconfigure(0, weight=1)

        # Controls
        controls_frame = ttk.LabelFrame(parent, text="Controls")
        controls_frame.grid(row=0, column=0, sticky="ew", padx=5, pady=5)
        ttk.Button(controls_frame, text="Run Health Check", command=self.handle_run_health_check).pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(controls_frame, text="Upload Document", command=self.handle_upload_file).pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(controls_frame, text="Force Save FAISS", command=self.handle_force_save_faiss).pack(fill=tk.X, padx=5, pady=5)

        # Profile
        profile_frame = ttk.LabelFrame(parent, text="Profile")
        profile_frame.grid(row=1, column=0, sticky="ew", padx=5, pady=5)
        ttk.Label(profile_frame, text="Name:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        self.profile_name_var = tk.StringVar()
        ttk.Entry(profile_frame, textvariable=self.profile_name_var).grid(row=0, column=1, sticky="ew", padx=5, pady=2)
        ttk.Label(profile_frame, text="Profession:").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        self.profile_prof_var = tk.StringVar()
        ttk.Entry(profile_frame, textvariable=self.profile_prof_var).grid(row=1, column=1, sticky="ew", padx=5, pady=2)
        ttk.Label(profile_frame, text="Interests:").grid(row=2, column=0, sticky="w", padx=5, pady=2)
        self.profile_int_var = tk.StringVar()
        ttk.Entry(profile_frame, textvariable=self.profile_int_var).grid(row=2, column=1, sticky="ew", padx=5, pady=2)
        ttk.Button(profile_frame, text="Save Profile", command=self.handle_save_profile).grid(row=3, columnspan=2, pady=5)
        self.refresh_profile_display()

        # Memory
        memory_frame = ttk.LabelFrame(parent, text="Memory & Retrieval")
        memory_frame.grid(row=2, column=0, sticky="ew", padx=5, pady=5)
        ttk.Label(memory_frame, text="Memory Window (Recent Messages):", wraplength=300).pack(fill=tk.X, padx=5)
        self.memory_slider = ttk.Scale(memory_frame, from_=0, to=50, orient=tk.HORIZONTAL, command=self.on_slider_change)
        self.memory_slider.set(self.state.memory_slider_val)
        self.memory_slider.pack(fill=tk.X, padx=5, pady=5)

        # Log Download
        log_frame = ttk.LabelFrame(parent, text="Logs")
        log_frame.grid(row=3, column=0, sticky="ew", padx=5, pady=5)
        ttk.Button(log_frame, text="Download Log File", command=self.handle_download_logs).pack(fill=tk.X, padx=5, pady=5)

    def _create_chat_widgets(self, parent):
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)

        # Chat display
        self.chat_display = scrolledtext.ScrolledText(parent, wrap=tk.WORD, state="disabled", font=("Helvetica", 11))
        self.chat_display.grid(row=0, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)
        self.chat_display.tag_config("user", foreground="blue", font=("Helvetica", 11, "bold"))
        self.chat_display.tag_config("assistant", foreground="black")

        # User input
        self.user_input_var = tk.StringVar()
        input_entry = ttk.Entry(parent, textvariable=self.user_input_var, font=("Helvetica", 11))
        input_entry.grid(row=1, column=0, sticky="ew", padx=5, pady=5)
        input_entry.bind("<Return>", lambda event: self.handle_send_message())
        
        send_button = ttk.Button(parent, text="Send", command=self.handle_send_message)
        send_button.grid(row=1, column=1, sticky="ew", padx=5, pady=5)

    def _create_dashboard_widgets(self, parent):
        parent.grid_columnconfigure(0, weight=1)

        # Health
        health_frame = ttk.LabelFrame(parent, text="Monitoring Dashboard")
        health_frame.grid(row=0, column=0, sticky="new", padx=5, pady=5)
        self.health_label = ttk.Label(health_frame, text="Run health check to see status.", justify=tk.LEFT, wraplength=400)
        self.health_label.pack(fill=tk.X, padx=5, pady=5)

        # Charts
        charts_notebook = ttk.Notebook(parent)
        charts_notebook.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)
        parent.grid_rowconfigure(1, weight=1)

        if MATPLOTLIB_AVAILABLE:
            self.charts = {}
            chart_specs = [
                ("Tokens", "LLM Tokens per Request", "requests", "tokens"),
                ("Latency", "LLM Latency (s)", "requests", "seconds"),
                ("Buffer", "Memory Buffer Size", "requests", "size"),
                ("Store", "FAISS Stores Size (MB)", "requests", "MB")
            ]
            for key, title, xlabel, ylabel in chart_specs:
                tab = ttk.Frame(charts_notebook)
                charts_notebook.add(tab, text=key)
                fig = Figure(figsize=(5, 2.5), dpi=100)
                ax = fig.add_subplot(111)
                ax.set_title(title, fontsize=10)
                ax.set_xlabel(xlabel, fontsize=8)
                ax.set_ylabel(ylabel, fontsize=8)
                ax.grid(True)
                fig.tight_layout()
                canvas = FigureCanvasTkAgg(fig, master=tab)
                canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=True)
                self.charts[key] = {'fig': fig, 'ax': ax, 'canvas': canvas}
        else:
            charts_notebook.add(ttk.Label(charts_notebook, text="Matplotlib not installed. Charts are disabled."), text="Charts")

        # Logs
        log_frame = ttk.LabelFrame(parent, text="Logs (tail)")
        log_frame.grid(row=2, column=0, sticky="sew", padx=5, pady=5)
        parent.grid_rowconfigure(2, weight=1) # Allow log to expand
        self.log_display = scrolledtext.ScrolledText(log_frame, wrap=tk.WORD, state="disabled", height=10, font=("Courier New", 9))
        self.log_display.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

    # --- Handlers and UI Update Methods ---
    
    def on_slider_change(self, value):
        self.state.memory_slider_val = int(float(value))
        if self.state.memory:
            self.state.memory.buffer_limit = self.state.memory_slider_val if self.state.memory_slider_val > 0 else CONFIG.SLIDING_WINDOW

    def handle_send_message(self):
        question = self.user_input_var.get().strip()
        if not question:
            return
        
        self.user_input_var.set("")
        self.add_message_to_chat("User", question)
        self.state.chat_history.append(("User", question))
        if self.state.memory: self.state.memory.add("user", question)
        
        # Run QA in a separate thread to not freeze the GUI
        threading.Thread(target=self._process_question_thread, args=(question,), daemon=True).start()
        self.after(100, self.check_response_queue)

    def _process_question_thread(self, question: str):
        if not self.state.qa or not self.state.vs or not self.state.profile_extractor:
            self.response_queue.put(("Error: QA engine not initialized.", {}))
            return

        try:
            extracted = self.state.profile_extractor.extract(question)
            if extracted:
                self.state.profile.update_from_dict(extracted)
                logger.info("Profile updated with keys: %s", list(extracted.keys()))
                self.after(0, self.refresh_profile_display) # Update GUI from main thread
        except Exception:
            logger.exception("Profile extraction error")
            
        docs = self.state.vs.search(question, k=CONFIG.TOP_K)
        answer, metrics = self.state.qa.answer(question, docs, memory_slider_val=self.state.memory_slider_val)
        self.response_queue.put((answer, metrics))

    def check_response_queue(self):
        try:
            answer, metrics = self.response_queue.get_nowait()
            self.add_message_to_chat("Assistant", answer)
            self.state.chat_history.append(("Assistant", answer))
            if self.state.memory: self.state.memory.add("assistant", answer)

            # Update metrics
            ts = time.time()
            self.state.metrics["tokens"].append(metrics.get("llm_tokens", 0))
            self.state.metrics["latencies"].append(metrics.get("latency", 0.0))
            self.state.metrics["buffer_sizes"].append(len(self.state.memory.buffer) if self.state.memory else 0)
            self.state.metrics["store_sizes"].append(folder_size_mb(Path(CONFIG.STORES_DIR)))
            self.state.metrics["timestamps"].append(ts)
            self.update_dashboard()

        except queue.Empty:
            self.after(100, self.check_response_queue)
    
    def handle_upload_file(self):
        filepath = filedialog.askopenfilename(
            title="Select a document",
            filetypes=[("All Files", "*.*"), ("PDF", "*.pdf"), ("Text", "*.txt"), ("Word", "*.docx"), ("JSON", "*.json"), ("Images", "*.png *.jpg *.jpeg")]
        )
        if not filepath:
            return

        path = Path(filepath)
        extractor = Extractor(Path(CONFIG.CACHE_DIR))
        text = ""
        suffix = path.suffix.lower()

        try:
            if suffix == ".pdf": text = extractor.from_pdf(str(path))
            elif suffix in {".txt", ".md"}: text = extractor.from_text(str(path))
            elif suffix == ".docx": text = extractor.from_docx(str(path))
            elif suffix == ".json": text = extractor.from_json(str(path))
            elif suffix in {".png", ".jpg", ".jpeg"}: text = extractor.from_image(str(path))
            else:
                messagebox.showwarning("Unsupported Type", f"Unsupported file type: {suffix}")
                return
            
            if text and self.state.indexer:
                fid = path.stem
                self.state.indexer.enqueue(text, fid)
                self.add_message_to_chat("System", f"Enqueued {path.name} for indexing as {fid}.")
            else:
                self.add_message_to_chat("System", f"Could not extract text from {path.name}.")
        except Exception as e:
            logger.exception("File processing failed")
            messagebox.showerror("File Error", f"Failed to process file: {e}")

    def handle_save_profile(self):
        if not self.state.profile: return
        self.state.profile.set("name", self.profile_name_var.get())
        self.state.profile.set("profession", self.profile_prof_var.get())
        interests = [s.strip() for s in self.profile_int_var.get().split(",") if s.strip()]
        self.state.profile.set("interests", interests)
        messagebox.showinfo("Profile", "Profile saved successfully.")

    def handle_run_health_check(self):
        if not self.state.qa: return
        self.state.last_health = check_health(self.state.qa.llm, self.state.vs.vertex_emb, self.state.vs.legal_emb, self.state.vs)
        health_text = "\n".join([f"- {k}: {v}" for k, v in self.state.last_health.items()])
        self.health_label.config(text=health_text)

    def handle_force_save_faiss(self):
        try:
            if self.state.vs:
                self.state.vs.save_all("global_docs")
                self.state.vs.save_all("chat_memory")
                messagebox.showinfo("Success", "FAISS stores saved.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save FAISS stores: {e}")
            
    def handle_download_logs(self):
        save_path = filedialog.asksaveasfilename(
            defaultextension=".log",
            initialfile=logfile.name,
            filetypes=[("Log Files", "*.log"), ("All files", "*.*")]
        )
        if save_path:
            try:
                import shutil
                shutil.copy(logfile, save_path)
                messagebox.showinfo("Success", f"Log file saved to {save_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Could not save log file: {e}")

    def add_message_to_chat(self, role, text):
        self.chat_display.config(state="normal")
        if role.lower() == "user":
            self.chat_display.insert(tk.END, "You: ", "user")
        elif role.lower() == "assistant":
            self.chat_display.insert(tk.END, "Assistant: ", "assistant_role")
        else: # System
            self.chat_display.insert(tk.END, f"{role}: ", ("system", "bold"))

        self.chat_display.insert(tk.END, f"{text}\n\n")
        self.chat_display.config(state="disabled")
        self.chat_display.see(tk.END)

    def update_log_display(self):
        log_tail = tail_log_lines(logfile, n=200)
        self.log_display.config(state="normal")
        self.log_display.delete(1.0, tk.END)
        self.log_display.insert(tk.END, log_tail)
        self.log_display.config(state="disabled")
        self.log_display.see(tk.END)
        self.after(5000, self.update_log_display) # Refresh every 5 seconds

    def update_dashboard(self):
        if not MATPLOTLIB_AVAILABLE: return
        metrics = self.state.metrics
        
        # Tokens
        ax = self.charts["Tokens"]["ax"]
        ax.clear()
        ax.plot(metrics["tokens"], marker='o', linestyle='-', markersize=4)
        ax.set_title("LLM Tokens per Request", fontsize=10)
        ax.grid(True)
        self.charts["Tokens"]["canvas"].draw()
        
        # Latency
        ax = self.charts["Latency"]["ax"]
        ax.clear()
        ax.plot(metrics["latencies"], marker='o', linestyle='-', markersize=4, color='orange')
        ax.set_title("LLM Latency (s)", fontsize=10)
        ax.grid(True)
        self.charts["Latency"]["canvas"].draw()

        # Buffer
        ax = self.charts["Buffer"]["ax"]
        ax.clear()
        ax.bar(range(len(metrics["buffer_sizes"])), metrics["buffer_sizes"], color='green')
        ax.set_title("Memory Buffer Size", fontsize=10)
        ax.grid(True)
        self.charts["Buffer"]["canvas"].draw()

        # Store Size
        ax = self.charts["Store"]["ax"]
        ax.clear()
        ax.plot(metrics["store_sizes"], marker='.', linestyle='-', markersize=4, color='red')
        ax.set_title("FAISS Stores Size (MB)", fontsize=10)
        ax.grid(True)
        self.charts["Store"]["canvas"].draw()

    def refresh_profile_display(self):
        if not self.state.profile: return
        profile_data = self.state.profile.all()
        self.profile_name_var.set(profile_data.get("name", ""))
        self.profile_prof_var.set(profile_data.get("profession", ""))
        interests = profile_data.get("interests", [])
        if isinstance(interests, list):
            self.profile_int_var.set(", ".join(interests))
        else:
            self.profile_int_var.set(interests or "")
            
    def on_closing(self):
        """Handle graceful shutdown of background processes."""
        logger.info("Shutdown sequence initiated...")
        if self.state.memory: self.state.memory.shutdown()
        if self.state.indexer: self.state.indexer.shutdown()
        if self.state.vs:
            self.state.vs.save_all("global_docs")
            self.state.vs.save_all("chat_memory")
        logger.info("Shutdown complete. Exiting.")
        self.destroy()

if __name__ == "__main__":
    TOKEN_MON = TokenMonitor()
    app = LegalAIGUI()
    app.mainloop()